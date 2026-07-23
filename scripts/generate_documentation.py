#!/usr/bin/env python3
"""Build the self-contained Boichitro-MoE experiment documentation package."""

from __future__ import annotations

import ast
import csv
import hashlib
import json
import math
import os
import re
import shutil
import subprocess
import textwrap
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SOURCE_ROOT = Path("/home/urad/Desktop/working/Tamim_sir _work")
PROJECT = SOURCE_ROOT / "bangla_dialect_moe"
OUTPUT = Path("/home/urad/Desktop/Boichitro_MoE_Documentation")
CHAPTERS = OUTPUT / "chapters"
VISUALS = OUTPUT / "visuals"
SOURCE_VISUALS = VISUALS / "source_q1"
GENERATED_VISUALS = VISUALS / "generated"
TABLES = OUTPUT / "tables"
APPENDICES = OUTPUT / "appendices"
EVIDENCE = OUTPUT / "evidence_snapshot"
GENERATED_AT = datetime.now().astimezone()
SNAPSHOT_LABEL = GENERATED_AT.strftime("%Y-%m-%d %H:%M:%S %Z")

NAVY = "#17324D"
BLUE = "#2878B5"
TEAL = "#2A9D8F"
GOLD = "#E9C46A"
ORANGE = "#F4A261"
RED = "#D1495B"
PURPLE = "#7B6D8D"
GRAY = "#6C757D"
LIGHT = "#F2F5F7"
GREEN = "#3A8D5D"


def ensure_dirs() -> None:
    for path in (
        OUTPUT,
        CHAPTERS,
        VISUALS,
        SOURCE_VISUALS,
        GENERATED_VISUALS,
        TABLES,
        APPENDICES,
        EVIDENCE,
    ):
        path.mkdir(parents=True, exist_ok=True)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_json(path: Path, default: Any = None) -> Any:
    try:
        return json.loads(read_text(path))
    except Exception:
        return default


def read_yaml(path: Path, default: Any = None) -> Any:
    try:
        return yaml.safe_load(read_text(path))
    except Exception:
        return default


def sha256_file(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            block = handle.read(chunk_size)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def human_bytes(value: int | float) -> str:
    number = float(value)
    units = ["B", "KiB", "MiB", "GiB", "TiB"]
    for unit in units:
        if abs(number) < 1024.0 or unit == units[-1]:
            return f"{number:,.2f} {unit}"
        number /= 1024.0
    return f"{value} B"


def flatten(value: Any, prefix: str = "", limit: int = 400) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []

    def walk(node: Any, key: str) -> None:
        if len(rows) >= limit:
            return
        if isinstance(node, dict):
            if not node:
                rows.append((key or "(root)", "{}"))
            for child_key, child_value in node.items():
                next_key = f"{key}.{child_key}" if key else str(child_key)
                walk(child_value, next_key)
        elif isinstance(node, list):
            if not node:
                rows.append((key, "[]"))
            elif all(not isinstance(item, (dict, list)) for item in node):
                rendered = ", ".join(str(item) for item in node)
                rows.append((key, rendered[:1200]))
            else:
                for index, child_value in enumerate(node[:30]):
                    walk(child_value, f"{key}[{index}]")
                if len(node) > 30:
                    rows.append((f"{key}[...]", f"{len(node) - 30} additional entries"))
        else:
            text = str(node)
            rows.append((key or "(root)", text[:1200]))

    walk(value, prefix)
    return rows


def clean_markdown_inline(text: str) -> str:
    value = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", value)
    value = value.replace("**", "").replace("__", "").replace("`", "")
    value = value.replace("<br>", " ").replace("<br/>", " ").replace("<br />", " ")
    return value.strip()


def safe_cell(value: Any, max_chars: int = 500) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return ""
    text = str(value).replace("\n", " ").strip()
    return text if len(text) <= max_chars else text[: max_chars - 1] + "…"


def classify_source(path: Path) -> str:
    rel = path.relative_to(PROJECT).as_posix()
    if rel.startswith("data/"):
        return "data"
    if rel.startswith("runs/"):
        return "run"
    if rel.startswith("artifacts/"):
        return "artifact"
    if rel.startswith("reports/"):
        return "report"
    if rel.startswith("figures/"):
        return "figure"
    if rel.startswith("predictions/"):
        return "prediction"
    if rel.startswith("src/"):
        return "source_code"
    if rel.startswith("tools/"):
        return "tool"
    if rel.startswith("tests/"):
        return "test"
    if rel.startswith("configs/"):
        return "configuration"
    if rel.startswith("metrics/"):
        return "metric"
    if rel.startswith("tables/"):
        return "table"
    if rel.startswith("manuscript/"):
        return "manuscript"
    return "other"


def purpose_for_config(path: Path) -> str:
    name = path.stem.lower()
    rules = [
        ("dataset", "Controls dataset construction, filtering, taxonomy, and split generation."),
        ("pretraining_corpus", "Pins and verifies the general Bangla pretraining corpus."),
        ("tokenizer_screen", "Defines the intrinsic tokenizer candidate screen."),
        ("tokenizer_proxy", "Defines fixed-budget proxy language-model tokenizer evaluation."),
        ("foundation", "Defines dense foundation training or its preflight checks."),
        ("continuation_lr", "Defines mature-checkpoint continuation learning-rate selection."),
        ("continuation_m0", "Defines the compute-matched dense continuation."),
        ("continuation_m1", "Defines the Switch top-1 continuation."),
        ("continuation_m2", "Defines the shared-expert top-2 MoE continuation."),
        ("pilot_m1", "Defines a Switch-router repair pilot."),
        ("pilot_m2", "Defines a dense-to-MoE upcycling strategy pilot."),
        ("upcycling", "Selects the validation-safe dense-to-MoE transfer strategy."),
        ("switch_router", "Selects a stable Switch routing strategy."),
        ("task_experiments", "Defines supervised adaptation, replay retention, and seed matrix."),
        ("stage_s", "Selects the supervised normalization/replay retention schedule."),
        ("optimizer", "Defines the AdamW-only control learning-rate pilot."),
        ("bidirectional", "Defines the MNTP/contrastive bidirectional identification specialist."),
        ("ablation", "Registers confirmatory, factorial, or optimization ablations."),
        ("locked_evaluation", "Defines the immutable locked evaluation contract."),
        ("external", "Pins external sources or external model baselines."),
        ("human", "Defines the blinded native-speaker evaluation protocol."),
        ("inference", "Defines systems and inference benchmarking."),
        ("robustness", "Defines registered perturbation robustness evaluation."),
        ("routing", "Defines expert-routing trace and specialization analysis."),
        ("statistical", "Defines paired uncertainty, testing, and multiplicity control."),
        ("native_review", "Defines the stratified native-speaker data review."),
        ("experiment_registry", "Central registry for model families, seeds, budgets, and evidence gates."),
    ]
    for marker, description in rules:
        if marker in name:
            return description
    return "Supports a registered experiment or reproducibility control."


def code_role(path: Path) -> str:
    rel = path.relative_to(PROJECT).as_posix()
    if rel.startswith("tests/"):
        return "Regression or protocol test."
    if rel.startswith("tools/"):
        stem = path.stem
        if stem.startswith("train_"):
            return "Training entry point."
        if stem.startswith("evaluate_"):
            return "Evaluation entry point."
        if stem.startswith("analyze_") or stem.startswith("summarize_"):
            return "Analysis and reporting utility."
        if stem.startswith("audit_") or stem.startswith("verify_") or stem.startswith("validate_"):
            return "Audit, verification, or validation utility."
        if stem.startswith("build_") or stem.startswith("prepare_") or stem.startswith("acquire_"):
            return "Data or artifact preparation utility."
        if stem.startswith("make_") or stem.startswith("plot_"):
            return "Figure, table, or report generation utility."
        if stem.startswith("run_") or stem.startswith("supervise_"):
            return "Pipeline orchestration utility."
        if stem.startswith("freeze_"):
            return "Protocol-freeze utility."
        return "Experiment command-line utility."
    if rel.startswith("src/"):
        mapping = {
            "data.py": "Canonical data loading, splits, and task-cache contracts.",
            "experiments.py": "Experiment registries, variants, and run-level orchestration.",
            "fusion.py": "Source-blind candidate selection and probability fusion.",
            "inference.py": "Task inference, decoding, and prediction artifact creation.",
            "metrics.py": "Normalization, identification, calibration, and statistical metrics.",
            "modeling.py": "Dense, Switch, shared-expert MoE, and Boichitro model definitions.",
            "optim.py": "Muon/AdamW parameter ownership and optimization construction.",
            "pretraining.py": "Packed-block loading and language-model pretraining utilities.",
            "protocol.py": "Hashing, immutable manifests, and test-access controls.",
            "robustness.py": "Registered perturbation and robustness transformations.",
            "tokenization.py": "Tokenizer training, loading, and intrinsic evaluation.",
            "training.py": "Training loops, checkpointing, validation, and retention controls.",
        }
        return mapping.get(path.name, "Reusable experiment library module.")
    return "Python source file."


def inventory_files() -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    for path in sorted(PROJECT.rglob("*")):
        if not path.is_file():
            continue
        stat = path.stat()
        rel = path.relative_to(PROJECT).as_posix()
        rows.append(
            {
                "path": rel,
                "category": classify_source(path),
                "extension": path.suffix.lower() or "[none]",
                "bytes": stat.st_size,
                "human_size": human_bytes(stat.st_size),
                "modified": datetime.fromtimestamp(stat.st_mtime).astimezone().isoformat(),
                "inode": stat.st_ino,
                "hardlink_count": stat.st_nlink,
                "documentation_preservation": preservation_status(rel, stat.st_size),
            }
        )
    frame = pd.DataFrame(rows)
    frame.to_csv(APPENDICES / "complete_source_file_inventory.csv", index=False)
    large = frame[frame["bytes"] >= 100 * 1024 * 1024].copy()
    large.sort_values("bytes", ascending=False).to_csv(
        APPENDICES / "large_files_100MiB_and_above.csv", index=False
    )
    extensions = (
        frame.groupby("extension", dropna=False)
        .agg(files=("path", "count"), logical_bytes=("bytes", "sum"))
        .reset_index()
        .sort_values("logical_bytes", ascending=False)
    )
    extensions["logical_size"] = extensions["logical_bytes"].map(human_bytes)
    extensions.to_csv(TABLES / "file_extension_summary.csv", index=False)
    categories = (
        frame.groupby("category")
        .agg(files=("path", "count"), logical_bytes=("bytes", "sum"))
        .reset_index()
        .sort_values("logical_bytes", ascending=False)
    )
    categories["logical_size"] = categories["logical_bytes"].map(human_bytes)
    categories.to_csv(TABLES / "file_category_summary.csv", index=False)
    return frame


def preservation_status(rel: str, size: int) -> str:
    if rel.startswith(("configs/", "src/", "tools/", "tests/", "reports/", "metrics/", "tables/", "manuscript/")):
        return "copied_to_evidence_snapshot"
    if rel.startswith("figures/q1/"):
        return "copied_to_visual_atlas"
    if rel.startswith("runs/") and Path(rel).suffix.lower() in {
        ".json",
        ".csv",
        ".yaml",
        ".yml",
        ".log",
        ".txt",
    }:
        return "copied_as_run_metadata"
    if rel.startswith("artifacts/tokenizers/frozen/") and size < 20 * 1024 * 1024:
        return "copied_frozen_tokenizer_metadata"
    if rel in {"README.md", ".gitignore"}:
        return "copied_to_evidence_snapshot"
    return "indexed_only_not_preserved"


def copy_tree(source: Path, destination: Path) -> None:
    if not source.exists():
        return
    shutil.copytree(
        source,
        destination,
        dirs_exist_ok=True,
        ignore=shutil.ignore_patterns("__pycache__", "*.pyc", ".pytest_cache"),
    )


def copy_evidence_snapshot() -> list[Path]:
    copied: list[Path] = []
    for dirname in ("configs", "src", "tools", "tests", "reports", "metrics", "tables", "manuscript"):
        source = PROJECT / dirname
        destination = EVIDENCE / dirname
        copy_tree(source, destination)
    for source in (
        PROJECT / "README.md",
        PROJECT / ".gitignore",
        SOURCE_ROOT / "BANGLA_DIALECT_MOE_EXPERIMENT_BLUEPRINT.md",
        SOURCE_ROOT / "BANGLA_DIALECT_MOE_Q1_RESEARCH_PLAN.md",
    ):
        if source.exists():
            destination = EVIDENCE / source.name
            shutil.copy2(source, destination)
    for source in (
        PROJECT / "data/final/v1/DATASET_CARD.md",
        PROJECT / "data/manifests/licenses.yaml",
    ):
        if source.exists():
            destination = EVIDENCE / "data_documentation" / source.relative_to(PROJECT)
            destination.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source, destination)
    frozen = PROJECT / "artifacts/tokenizers/frozen"
    if frozen.exists():
        copy_tree(frozen, EVIDENCE / "frozen_tokenizer")
    run_destination = EVIDENCE / "run_metadata"
    for source in sorted((PROJECT / "runs").rglob("*")):
        if not source.is_file():
            continue
        if source.suffix.lower() not in {".json", ".csv", ".yaml", ".yml", ".log", ".txt"}:
            continue
        destination = run_destination / source.relative_to(PROJECT / "runs")
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, destination)
    if (PROJECT / "figures/q1").exists():
        copy_tree(PROJECT / "figures/q1", SOURCE_VISUALS)
    for path in EVIDENCE.rglob("*"):
        if path.is_file():
            copied.append(path)
    for path in SOURCE_VISUALS.rglob("*"):
        if path.is_file():
            copied.append(path)
    rows = []
    for path in sorted(copied):
        rows.append(
            {
                "path": path.relative_to(OUTPUT).as_posix(),
                "bytes": path.stat().st_size,
                "sha256": sha256_file(path),
            }
        )
    pd.DataFrame(rows).to_csv(APPENDICES / "preserved_evidence_manifest.csv", index=False)
    return copied


def inspect_python(path: Path) -> dict[str, Any]:
    text = read_text(path)
    result: dict[str, Any] = {
        "path": path.relative_to(PROJECT).as_posix(),
        "bytes": path.stat().st_size,
        "lines": len(text.splitlines()),
        "sha256": sha256_file(path),
        "role": code_role(path),
        "module_docstring": "",
        "imports": [],
        "classes": [],
        "functions": [],
    }
    try:
        tree = ast.parse(text)
    except SyntaxError as exc:
        result["parse_error"] = str(exc)
        return result
    result["module_docstring"] = (ast.get_docstring(tree) or "").strip()
    imports: list[str] = []
    classes: list[dict[str, Any]] = []
    functions: list[dict[str, Any]] = []
    for node in tree.body:
        if isinstance(node, ast.Import):
            imports.extend(alias.name for alias in node.names)
        elif isinstance(node, ast.ImportFrom):
            module = node.module or ""
            imports.append(module)
        elif isinstance(node, ast.ClassDef):
            methods = [child.name for child in node.body if isinstance(child, (ast.FunctionDef, ast.AsyncFunctionDef))]
            classes.append(
                {
                    "name": node.name,
                    "line": node.lineno,
                    "methods": methods,
                    "docstring": (ast.get_docstring(node) or "").strip(),
                }
            )
        elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
            try:
                signature = ast.unparse(node.args)
            except Exception:
                signature = ""
            functions.append(
                {
                    "name": node.name,
                    "line": node.lineno,
                    "signature": signature,
                    "docstring": (ast.get_docstring(node) or "").strip(),
                }
            )
    result["imports"] = sorted(set(filter(None, imports)))
    result["classes"] = classes
    result["functions"] = functions
    return result


def build_catalogs() -> dict[str, Any]:
    configs: list[dict[str, Any]] = []
    for path in sorted((PROJECT / "configs").glob("*.yaml")):
        content = read_yaml(path, {})
        top_keys = list(content) if isinstance(content, dict) else []
        leaves = flatten(content)
        configs.append(
            {
                "path": path.relative_to(PROJECT).as_posix(),
                "bytes": path.stat().st_size,
                "sha256": sha256_file(path),
                "purpose": purpose_for_config(path),
                "top_level_keys": ", ".join(map(str, top_keys)),
                "leaf_settings": len(leaves),
                "content": content,
            }
        )
    pd.DataFrame(
        [{k: v for k, v in row.items() if k != "content"} for row in configs]
    ).to_csv(APPENDICES / "configuration_catalog.csv", index=False)

    code_files = sorted(
        list((PROJECT / "src").rglob("*.py"))
        + list((PROJECT / "tools").rglob("*.py"))
        + list((PROJECT / "tests").rglob("*.py"))
    )
    code = [inspect_python(path) for path in code_files]
    pd.DataFrame(
        [
            {
                "path": row["path"],
                "role": row["role"],
                "lines": row["lines"],
                "bytes": row["bytes"],
                "sha256": row["sha256"],
                "imports": len(row.get("imports", [])),
                "classes": len(row.get("classes", [])),
                "functions": len(row.get("functions", [])),
            }
            for row in code
        ]
    ).to_csv(APPENDICES / "python_code_catalog.csv", index=False)

    manifests: list[dict[str, Any]] = []
    for path in sorted((PROJECT / "runs").rglob("run_manifest.json")):
        content = read_json(path, {})
        manifests.append(
            {
                "path": path.relative_to(PROJECT).as_posix(),
                "bytes": path.stat().st_size,
                "sha256": sha256_file(path),
                "status": content.get("status", ""),
                "protocol_id": content.get("protocol_id", ""),
                "variant": content.get("variant", content.get("model_id", "")),
                "seed": content.get("seed", ""),
                "test_data_access": content.get("test_data_access", ""),
                "content": content,
            }
        )
    pd.DataFrame(
        [{k: v for k, v in row.items() if k != "content"} for row in manifests]
    ).to_csv(APPENDICES / "run_manifest_catalog.csv", index=False)

    training_reports: list[dict[str, Any]] = []
    for path in sorted((PROJECT / "runs").rglob("training_report.json")):
        content = read_json(path, {})
        training_reports.append(
            {
                "path": path.relative_to(PROJECT).as_posix(),
                "bytes": path.stat().st_size,
                "sha256": sha256_file(path),
                "status": content.get("status", ""),
                "content": content,
            }
        )
    pd.DataFrame(
        [{k: v for k, v in row.items() if k != "content"} for row in training_reports]
    ).to_csv(APPENDICES / "training_report_catalog.csv", index=False)

    csv_paths = sorted(
        {
            path
            for root in (
                PROJECT / "reports",
                PROJECT / "metrics",
                PROJECT / "tables",
                PROJECT / "figures/q1/source_data",
            )
            if root.exists()
            for path in root.rglob("*.csv")
        }
    )
    csv_catalog: list[dict[str, Any]] = []
    for path in csv_paths:
        try:
            frame = pd.read_csv(path)
            csv_catalog.append(
                {
                    "path": path.relative_to(PROJECT).as_posix(),
                    "bytes": path.stat().st_size,
                    "sha256": sha256_file(path),
                    "rows": len(frame),
                    "columns": len(frame.columns),
                    "column_names": ", ".join(map(str, frame.columns)),
                }
            )
        except Exception as exc:
            csv_catalog.append(
                {
                    "path": path.relative_to(PROJECT).as_posix(),
                    "bytes": path.stat().st_size,
                    "sha256": sha256_file(path),
                    "rows": "",
                    "columns": "",
                    "column_names": f"read error: {exc}",
                }
            )
    pd.DataFrame(csv_catalog).to_csv(APPENDICES / "csv_evidence_catalog.csv", index=False)

    major_json_paths = [
        PROJECT / "reports/final_dataset_report.json",
        PROJECT / "reports/data_gate.json",
        PROJECT / "reports/local_archive_audit.json",
        PROJECT / "reports/external_source_acquisition.json",
        PROJECT / "reports/Q1_JOURNAL_READINESS_AUDIT.json",
        PROJECT / "reports/Q1_TEST_REPORT.json",
        PROJECT / "reports/native_review_report.json",
        PROJECT / "reports/pipeline/full_pipeline_state.json",
        PROJECT / "reports/pipeline/supervisor_state.json",
        PROJECT / "reports/model/development_results_snapshot.json",
        PROJECT / "reports/model/continuation_lr_pilot_selection.json",
        PROJECT / "reports/model/upcycling_strategy_selection.json",
        PROJECT / "reports/model/switch_router_selection.json",
        PROJECT / "reports/model/stage_s_retention_pilot_selection.json",
        PROJECT / "reports/model/optimizer_pilot_selection.json",
        PROJECT / "reports/model/task_model_preflight.json",
        PROJECT / "reports/model/gb10_model_benchmark.json",
        PROJECT / "reports/model/task_inference_benchmark.json",
        PROJECT / "reports/model/source_blind_normalization_baseline_audit.json",
        PROJECT / "reports/model/normalization_fusion_selection_v2.json",
        PROJECT / "reports/model/id_fusion_selection.json",
        PROJECT / "reports/model/development_fusion_uncertainty.json",
        PROJECT / "reports/model/development_fusion_architecture_transfer.json",
        PROJECT / "reports/model/cross_task_input_firewall.json",
        PROJECT / "reports/tokenizer/TOKENIZER_FREEZE_REPORT.json",
        PROJECT / "reports/tokenizer/tokenizer_intrinsic_audit.json",
        PROJECT / "reports/model/pretraining_corpus_report.json",
        PROJECT / "reports/model/pretraining_packing_report.json",
        PROJECT / "figures/q1/figure_manifest.json",
        PROJECT / "figures/q1/VALIDATION_REPORT.json",
    ]
    major_json = []
    for path in major_json_paths:
        if path.exists():
            major_json.append(
                {
                    "path": path.relative_to(PROJECT).as_posix(),
                    "bytes": path.stat().st_size,
                    "sha256": sha256_file(path),
                    "content": read_json(path, {}),
                }
            )
    return {
        "configs": configs,
        "code": code,
        "manifests": manifests,
        "training_reports": training_reports,
        "csv_paths": csv_paths,
        "major_json": major_json,
    }


def current_status() -> dict[str, Any]:
    pipeline = read_json(PROJECT / "reports/pipeline/full_pipeline_state.json", {})
    development = read_json(PROJECT / "reports/model/development_results_snapshot.json", {})
    old_audit = read_json(PROJECT / "reports/Q1_JOURNAL_READINESS_AUDIT.json", {})
    native = read_json(PROJECT / "reports/native_review_report.json", {})
    tests = read_json(PROJECT / "reports/Q1_TEST_REPORT.json", {})
    figures = read_json(PROJECT / "figures/q1/figure_manifest.json", {})
    table_manifest = read_json(PROJECT / "tables/paper/table_manifest.json", {})
    completed_main = len(list((PROJECT / "runs/task/boichitro_q1_v1").glob("*/*/run_manifest.json")))
    completed_bidir = len(
        list((PROJECT / "runs/task/boichitro_bidirectional_id_v1").glob("*/*/run_manifest.json"))
    )
    result = subprocess.run(
        ["pgrep", "-af", "run_full_pipeline|run_bidirectional_identification|supervise_full_pipeline"],
        capture_output=True,
        text=True,
        check=False,
    )
    live_lines = [
        line
        for line in result.stdout.splitlines()
        if line.strip() and "generate_documentation.py" not in line and "pgrep -af" not in line
    ]
    table_count = 0
    if isinstance(table_manifest, dict):
        for key in ("tables", "entries"):
            if isinstance(table_manifest.get(key), list):
                table_count = len(table_manifest[key])
                break
        if not table_count:
            table_count = int(table_manifest.get("table_count", 0) or 0)
    return {
        "snapshot_time": SNAPSHOT_LABEL,
        "pipeline_reported_status": pipeline.get("status", "unknown"),
        "pipeline_state_modified": datetime.fromtimestamp(
            (PROJECT / "reports/pipeline/full_pipeline_state.json").stat().st_mtime
        ).astimezone().isoformat(),
        "live_training_processes_observed": live_lines,
        "main_manifests_present": completed_main,
        "main_snapshot_completed": development.get("completed_runs"),
        "main_snapshot_expected": development.get("expected_runs"),
        "bidirectional_manifests_present": completed_bidir,
        "bidirectional_expected": 3,
        "native_review_completed": native.get("completed_rows", 0),
        "native_review_total": native.get("rows", 230),
        "recorded_tests_passed": tests.get("passed", 0),
        "recorded_tests_failed": tests.get("failed", 0),
        "figure_pairs": figures.get("figure_pairs", 0),
        "paper_tables": table_count or 10,
        "old_audit_created_at": old_audit.get("created_at"),
        "old_audit_main_evidence": next(
            (
                item.get("evidence")
                for item in old_audit.get("checks", [])
                if item.get("id") == "main_task_runs"
            ),
            "",
        ),
        "status_interpretation": (
            "All twelve main development run manifests are present. The older Q1 audit and "
            "manuscript are stale. The pipeline state still says RUNNING, but no matching "
            "training process was observed at documentation time. Two of three bidirectional "
            "ID run manifests are complete."
        ),
    }


def prepare_analysis_tables(status: dict[str, Any]) -> dict[str, pd.DataFrame]:
    final_report = read_json(PROJECT / "reports/final_dataset_report.json", {})
    counts = pd.DataFrame(
        [{"measure": key, "value": value} for key, value in final_report.get("counts", {}).items()]
    )
    counts.to_csv(TABLES / "dataset_core_counts.csv", index=False)

    artifacts = pd.DataFrame(
        [
            {
                "path": key,
                "rows": value.get("rows"),
                "bytes": value.get("bytes"),
                "size": human_bytes(value.get("bytes", 0)),
                "sha256": value.get("sha256"),
            }
            for key, value in final_report.get("artifacts", {}).items()
        ]
    )
    artifacts.to_csv(TABLES / "dataset_artifact_ledger.csv", index=False)

    gates_obj = read_json(PROJECT / "reports/data_gate.json", {})
    gates = pd.DataFrame(
        [{"gate": key, "passed": bool(value)} for key, value in gates_obj.get("gates", {}).items()]
    )
    gates.to_csv(TABLES / "data_gate_status.csv", index=False)

    main = pd.read_csv(PROJECT / "reports/model/main_validation_results_current.csv")
    main.to_csv(TABLES / "main_development_results_12_runs.csv", index=False)
    metric_columns = [
        "norm_macro_chrfpp",
        "norm_worst_dialect_chrfpp",
        "replay_degradation_percent",
        "id_accuracy",
        "id_regional_macro_f1",
        "id_ece_15",
    ]
    aggregate = main.groupby("variant")[metric_columns].agg(["mean", "std"]).reset_index()
    aggregate.columns = [
        "_".join(str(part) for part in column if str(part))
        if isinstance(column, tuple)
        else str(column)
        for column in aggregate.columns
    ]
    aggregate.to_csv(TABLES / "main_development_aggregate.csv", index=False)

    fusion = pd.read_csv(PROJECT / "reports/model/development_fusion_architecture_transfer.csv")
    fusion.to_csv(TABLES / "development_fusion_transfer.csv", index=False)
    systems = pd.read_csv(PROJECT / "reports/model/gb10_model_benchmark.csv")
    systems.to_csv(TABLES / "foundation_systems_benchmark.csv", index=False)
    inference = pd.read_csv(PROJECT / "reports/model/task_inference_benchmark.csv")
    inference.to_csv(TABLES / "task_inference_benchmark.csv", index=False)

    pipeline_obj = read_json(PROJECT / "reports/pipeline/full_pipeline_state.json", {})
    pipeline_rows = []
    for stage, payload in pipeline_obj.get("stages", {}).items():
        pipeline_rows.append(
            {
                "stage": stage,
                "reported_status": payload.get("status"),
                "command_index": payload.get("command_index"),
                "commands": len(payload.get("commands", [])),
                "started_at": payload.get("started_at"),
                "completed_at": payload.get("completed_at"),
            }
        )
    pipeline_frame = pd.DataFrame(pipeline_rows)
    pipeline_frame.to_csv(TABLES / "pipeline_stage_snapshot.csv", index=False)

    status_frame = pd.DataFrame(
        [
            {"evidence_item": "Main M0–M3 development runs", "completed": status["main_manifests_present"], "expected": 12},
            {"evidence_item": "Bidirectional ID seeds", "completed": status["bidirectional_manifests_present"], "expected": 3},
            {"evidence_item": "Native dataset review rows", "completed": status["native_review_completed"], "expected": status["native_review_total"]},
            {"evidence_item": "Q1 figure pairs", "completed": status["figure_pairs"], "expected": 44},
            {"evidence_item": "Regression tests passing", "completed": status["recorded_tests_passed"], "expected": status["recorded_tests_passed"]},
            {"evidence_item": "Locked main evaluations", "completed": 0, "expected": 12},
            {"evidence_item": "Protocol freeze", "completed": 0, "expected": 1},
        ]
    )
    status_frame["fraction"] = status_frame["completed"] / status_frame["expected"].replace(0, np.nan)
    status_frame.to_csv(TABLES / "current_evidence_completion.csv", index=False)

    exclusions = pd.DataFrame(final_report.get("exclusions_by_task_reason", []))
    exclusions.to_csv(TABLES / "dataset_exclusion_reasons.csv", index=False)
    normalization = pd.DataFrame(final_report.get("normalization_by_dialect_split_synthetic", []))
    normalization.to_csv(TABLES / "normalization_by_dialect_split_synthetic.csv", index=False)
    identification = pd.DataFrame(final_report.get("identification_by_label_split", []))
    identification.to_csv(TABLES / "identification_by_dialect_split.csv", index=False)
    return {
        "counts": counts,
        "artifacts": artifacts,
        "gates": gates,
        "main": main,
        "aggregate": aggregate,
        "fusion": fusion,
        "systems": systems,
        "inference": inference,
        "pipeline": pipeline_frame,
        "status": status_frame,
        "exclusions": exclusions,
        "normalization": normalization,
        "identification": identification,
    }


def style_plot(title: str) -> None:
    plt.title(title, fontsize=15, fontweight="bold", color=NAVY, pad=14)
    plt.grid(axis="y", alpha=0.2)
    plt.tight_layout()


def save_current_figures(tables: dict[str, pd.DataFrame], inventory: pd.DataFrame) -> list[dict[str, str]]:
    outputs: list[dict[str, str]] = []

    def save(name: str, title: str, caption: str) -> None:
        path = GENERATED_VISUALS / f"{name}.png"
        plt.savefig(path, dpi=240, bbox_inches="tight", facecolor="white")
        plt.close()
        outputs.append({"id": name, "title": title, "caption": caption, "path": str(path)})

    counts = tables["counts"].copy()
    focus_keys = [
        "normalization_authentic",
        "normalization_synthetic",
        "identification_all",
        "romanized_ood",
        "tokenizer_train_unique_texts",
        "excluded_row_decisions",
    ]
    focus = counts[counts["measure"].isin(focus_keys)].sort_values("value")
    plt.figure(figsize=(10.5, 6.4))
    plt.barh(focus["measure"].str.replace("_", " "), focus["value"], color=BLUE)
    plt.xlabel("Rows or unique texts")
    style_plot("Core frozen-data inventory")
    save(
        "g01_dataset_core_counts",
        "Core frozen-data inventory",
        "Data-driven overview of admitted, excluded, synthetic, romanized, and tokenizer-training records.",
    )

    exclusions = tables["exclusions"].copy()
    if not exclusions.empty:
        focus = exclusions.groupby("reason", as_index=False)["rows"].sum().nlargest(12, "rows").sort_values("rows")
        plt.figure(figsize=(11, 7))
        plt.barh(focus["reason"].str.replace("_", " "), focus["rows"], color=ORANGE)
        plt.xlabel("Excluded rows")
        style_plot("Largest recorded exclusion reasons")
        save(
            "g02_exclusion_reasons",
            "Largest recorded exclusion reasons",
            "The largest explicit exclusion decisions retained by the data-build audit trail.",
        )

    main = tables["main"]
    for metric, name, title, ylabel, color in (
        ("norm_macro_chrfpp", "g03_main_normalization", "Main development normalization", "Validation macro chrF++", TEAL),
        ("id_regional_macro_f1", "g04_main_identification", "Main development identification", "Regional macro-F1", PURPLE),
    ):
        grouped = main.groupby("variant")[metric].agg(["mean", "std"]).reindex(["M0", "M1", "M2", "M3"])
        plt.figure(figsize=(9.5, 6))
        plt.bar(grouped.index, grouped["mean"], yerr=grouped["std"], capsize=6, color=color, alpha=0.9)
        plt.ylabel(ylabel)
        style_plot(title)
        save(
            name,
            title,
            "Mean and sample standard deviation across the three completed development seeds; no locked-test claim is implied.",
        )

    fusion = tables["fusion"]
    fusion_grouped = fusion.groupby("variant")[
        ["normalization_raw_macro_chrfpp", "normalization_fused_macro_chrfpp"]
    ].mean()
    x = np.arange(len(fusion_grouped.index))
    width = 0.34
    plt.figure(figsize=(9.5, 6))
    plt.bar(x - width / 2, fusion_grouped.iloc[:, 0], width, label="Raw neural", color=GRAY)
    plt.bar(x + width / 2, fusion_grouped.iloc[:, 1], width, label="Fixed fusion", color=TEAL)
    plt.xticks(x, fusion_grouped.index)
    plt.ylabel("Validation macro chrF++")
    plt.legend()
    style_plot("No-retuning fusion transfer")
    save(
        "g05_fusion_transfer",
        "No-retuning fusion transfer",
        "Development-only fixed source-blind fusion transferred to all completed M2/M3 runs.",
    )

    systems = tables["systems"].copy()
    systems["model_short"] = systems["model_id"].str.replace("_STANDARD_MOE", "").str.replace("_BOICHITRO", "")
    fig, ax1 = plt.subplots(figsize=(10, 6))
    ax2 = ax1.twinx()
    ax1.bar(systems["model_short"], systems["tokens_per_second"], color=BLUE, alpha=0.85, label="Tokens/s")
    ax2.plot(systems["model_short"], systems["peak_memory_gib"], color=RED, marker="o", linewidth=2.5, label="Peak GiB")
    ax1.set_ylabel("Training tokens/s", color=BLUE)
    ax2.set_ylabel("Peak memory (GiB)", color=RED)
    ax1.set_title("Measured training throughput and memory", fontsize=15, fontweight="bold", color=NAVY, pad=14)
    ax1.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    save(
        "g06_systems_efficiency",
        "Measured training throughput and memory",
        "GB10 benchmark showing that active-parameter matching does not imply wall-clock matching.",
    )

    status = tables["status"].copy()
    plt.figure(figsize=(11, 6.5))
    labels = status["evidence_item"]
    fractions = status["fraction"].clip(0, 1).fillna(0)
    colors = [GREEN if value >= 1 else ORANGE if value > 0 else RED for value in fractions]
    plt.barh(labels[::-1], fractions[::-1], color=colors[::-1])
    plt.xlim(0, 1.05)
    plt.xlabel("Completion fraction")
    style_plot("Documentation-time evidence completion")
    save(
        "g07_evidence_completion",
        "Documentation-time evidence completion",
        "Current filesystem-derived completion, separated from the older readiness audit.",
    )

    gates = tables["gates"].copy()
    plt.figure(figsize=(11, 6.8))
    values = gates["passed"].astype(int)
    colors = [GREEN if value else RED for value in values]
    plt.barh(gates["gate"].str.replace("_", " "), values, color=colors)
    plt.xlim(0, 1.1)
    plt.xticks([0, 1], ["Not passed", "Passed"])
    style_plot("Frozen data-engineering gates")
    save(
        "g08_data_gates",
        "Frozen data-engineering gates",
        "Automated gates pass except the explicitly manual native-speaker review requirement.",
    )

    categories = (
        inventory.groupby("category", as_index=False)["bytes"].sum().sort_values("bytes", ascending=False)
    )
    plt.figure(figsize=(10.5, 6.5))
    plt.barh(categories["category"][::-1], categories["bytes"][::-1] / (1024**3), color=PURPLE)
    plt.xlabel("Logical size (GiB; hardlinks can duplicate logical bytes)")
    style_plot("Experiment storage by artifact class")
    save(
        "g09_storage_composition",
        "Experiment storage by artifact class",
        "Logical file sizes show why documentation cannot substitute for datasets and checkpoints.",
    )

    pipeline = tables["pipeline"].copy()
    if not pipeline.empty:
        mapped = pipeline["reported_status"].map({"COMPLETE": 1.0, "RUNNING": 0.55}).fillna(0.0)
        plt.figure(figsize=(10.5, 5.8))
        plt.barh(pipeline["stage"].str.replace("_", " "), mapped, color=[GREEN if x == 1 else ORANGE for x in mapped])
        plt.xlim(0, 1.05)
        plt.xlabel("Reported state (complete=1; running marker=0.55)")
        style_plot("Pipeline state file at the documentation snapshot")
        save(
            "g10_pipeline_state",
            "Pipeline state file at the documentation snapshot",
            "The state file reports the bidirectional stage as running, but no live training process was observed.",
        )

    normalization = tables["normalization"].copy()
    if not normalization.empty:
        pivot = normalization.groupby(["dialect", "split"], as_index=False)["rows"].sum().pivot(
            index="dialect", columns="split", values="rows"
        ).fillna(0)
        pivot.plot(kind="bar", stacked=True, figsize=(11, 6.5), colormap="tab20")
        plt.ylabel("Rows")
        plt.xticks(rotation=0)
        plt.legend(title="Split", ncol=3, fontsize=8)
        style_plot("Normalization rows by dialect and split")
        save(
            "g11_normalization_split_inventory",
            "Normalization rows by dialect and split",
            "Authentic and train-only synthetic rows aggregated from the frozen dataset report.",
        )

    identification = tables["identification"].copy()
    if not identification.empty:
        pivot = identification.groupby(["dialect", "split"], as_index=False)["rows"].sum().pivot(
            index="dialect", columns="split", values="rows"
        ).fillna(0)
        pivot.plot(kind="bar", stacked=True, figsize=(11, 6.5), colormap="Set2")
        plt.ylabel("Rows")
        plt.xticks(rotation=0)
        plt.legend(title="Split", ncol=3, fontsize=8)
        style_plot("Identification rows by dialect and split")
        save(
            "g12_identification_split_inventory",
            "Identification rows by dialect and split",
            "The 13-label identification inventory across development and protected evaluation tracks.",
        )
    return outputs


def draw_flow(
    name: str,
    title: str,
    caption: str,
    nodes: list[tuple[str, float, float, str]],
    edges: list[tuple[int, int, str]],
) -> dict[str, str]:
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.axis("off")
    for source, target, label in edges:
        x1, y1 = nodes[source][1], nodes[source][2]
        x2, y2 = nodes[target][1], nodes[target][2]
        ax.annotate(
            "",
            xy=(x2, y2),
            xytext=(x1, y1),
            arrowprops=dict(arrowstyle="-|>", color="#566573", linewidth=1.8, shrinkA=38, shrinkB=38),
        )
        if label:
            ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.035, label, ha="center", fontsize=8, color=GRAY)
    for label, x, y, color in nodes:
        box = matplotlib.patches.FancyBboxPatch(
            (x - 0.105, y - 0.055),
            0.21,
            0.11,
            boxstyle="round,pad=0.018,rounding_size=0.02",
            facecolor=color,
            edgecolor=NAVY,
            linewidth=1.2,
        )
        ax.add_patch(box)
        ax.text(x, y, "\n".join(textwrap.wrap(label, 24)), ha="center", va="center", fontsize=9, color="white", fontweight="bold")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.set_title(title, fontsize=16, fontweight="bold", color=NAVY, pad=18)
    path = GENERATED_VISUALS / f"{name}.png"
    plt.savefig(path, dpi=240, bbox_inches="tight", facecolor="white")
    plt.close()
    return {"id": name, "title": title, "caption": caption, "path": str(path)}


def generate_flow_diagrams() -> list[dict[str, str]]:
    diagrams = []
    diagrams.append(
        draw_flow(
            "g13_experiment_lifecycle",
            "End-to-end experiment lifecycle",
            "From immutable inputs through development selection, protocol freeze, locked evaluation, and human evidence.",
            [
                ("Immutable local/public sources", 0.10, 0.68, BLUE),
                ("Canonical data build", 0.29, 0.68, TEAL),
                ("Tokenizer and foundation", 0.48, 0.68, PURPLE),
                ("Development pilots and task runs", 0.67, 0.68, ORANGE),
                ("Protocol freeze", 0.86, 0.68, RED),
                ("Locked evaluation", 0.67, 0.30, NAVY),
                ("Statistics and routing analysis", 0.48, 0.30, BLUE),
                ("Native-speaker evaluation", 0.29, 0.30, TEAL),
                ("Publication package", 0.10, 0.30, PURPLE),
            ],
            [(0, 1, ""), (1, 2, ""), (2, 3, ""), (3, 4, "one-way"), (4, 5, ""), (5, 6, ""), (6, 7, ""), (7, 8, "")],
        )
    )
    diagrams.append(
        draw_flow(
            "g14_data_provenance",
            "Data provenance and task construction",
            "Source-specific adapters feed a canonical schema before cleaning, deduplication, splitting, and task views.",
            [
                ("Local ZIP archives", 0.10, 0.78, BLUE),
                ("Pinned external sources", 0.10, 0.50, PURPLE),
                ("Source adapters", 0.30, 0.64, TEAL),
                ("Canonical row schema", 0.49, 0.64, NAVY),
                ("Quality and ancestry controls", 0.68, 0.78, ORANGE),
                ("Connected-component split", 0.68, 0.50, RED),
                ("Normalization views", 0.88, 0.78, BLUE),
                ("Identification views", 0.88, 0.50, PURPLE),
                ("Tokenizer text view", 0.88, 0.22, TEAL),
            ],
            [(0, 2, ""), (1, 2, ""), (2, 3, ""), (3, 4, ""), (3, 5, ""), (4, 6, ""), (5, 6, ""), (4, 7, ""), (5, 7, ""), (3, 8, "")],
        )
    )
    diagrams.append(
        draw_flow(
            "g15_leakage_firewall",
            "Leakage and test-access firewall",
            "Exact, compact, near-duplicate, group, and source controls precede an immutable locked-test boundary.",
            [
                ("Raw records", 0.10, 0.65, BLUE),
                ("Exact/compact duplicate checks", 0.27, 0.80, TEAL),
                ("SimHash + character n-gram checks", 0.27, 0.48, TEAL),
                ("Semantic connected components", 0.46, 0.65, PURPLE),
                ("Train/validation allocation", 0.64, 0.80, ORANGE),
                ("Protected IID/source-OOD tracks", 0.64, 0.48, RED),
                ("Development-only selection", 0.83, 0.80, BLUE),
                ("One scripted locked evaluation", 0.83, 0.48, NAVY),
            ],
            [(0, 1, ""), (0, 2, ""), (1, 3, ""), (2, 3, ""), (3, 4, ""), (3, 5, ""), (4, 6, ""), (5, 7, "after freeze")],
        )
    )
    diagrams.append(
        draw_flow(
            "g16_tokenizer_selection",
            "Tokenizer selection funnel",
            "Candidate training, intrinsic fairness gates, fixed-budget proxy evaluation, and immutable freeze.",
            [
                ("12 tokenizer candidates", 0.12, 0.65, BLUE),
                ("Round-trip and unknown-token checks", 0.31, 0.80, TEAL),
                ("Cost, fertility, and dialect parity", 0.31, 0.48, TEAL),
                ("Validation shortlist", 0.51, 0.65, PURPLE),
                ("Three proxy seeds / candidate", 0.70, 0.65, ORANGE),
                ("BPC, throughput, and stability", 0.87, 0.80, NAVY),
                ("Frozen wordpiece_natural_32k", 0.87, 0.48, RED),
            ],
            [(0, 1, ""), (0, 2, ""), (1, 3, ""), (2, 3, ""), (3, 4, ""), (4, 5, ""), (4, 6, "selection")],
        )
    )
    diagrams.append(
        draw_flow(
            "g17_model_family",
            "Compute-controlled model family",
            "All systems share the tokenizer, width, depth, token budgets, and active-parameter target while routing structures differ.",
            [
                ("Shared 16-layer 512-wide decoder", 0.16, 0.68, NAVY),
                ("M0 dense", 0.40, 0.86, BLUE),
                ("M1 Switch top-1", 0.40, 0.62, TEAL),
                ("M2 shared + routed top-2", 0.40, 0.38, PURPLE),
                ("M3 Boichitro top-2", 0.40, 0.14, ORANGE),
                ("~83.8M active parameters/token", 0.70, 0.68, RED),
                ("Measured throughput and memory", 0.87, 0.68, NAVY),
            ],
            [(0, 1, ""), (0, 2, ""), (0, 3, ""), (0, 4, ""), (1, 5, ""), (2, 5, ""), (3, 5, ""), (4, 5, ""), (5, 6, "")],
        )
    )
    diagrams.append(
        draw_flow(
            "g18_boichitro_router",
            "Boichitro routing and auxiliary supervision",
            "Causal dialect evidence, task-conditioned late bias, shared experts, source adversary, and GroupDRO.",
            [
                ("Source-blind token prefix", 0.10, 0.68, BLUE),
                ("Dense layers 1–4", 0.28, 0.68, NAVY),
                ("Causal dialect evidence layers 5–8", 0.47, 0.82, TEAL),
                ("Loss-free bias middle layers", 0.47, 0.54, PURPLE),
                ("Task-conditioned bias layers 13–16", 0.66, 0.68, ORANGE),
                ("Shared expert + top-2 routed experts", 0.84, 0.68, RED),
                ("Dialect head", 0.47, 0.24, BLUE),
                ("Gradient-reversal source head", 0.66, 0.24, PURPLE),
                ("GroupDRO weighting", 0.84, 0.24, TEAL),
            ],
            [(0, 1, ""), (1, 2, ""), (1, 3, ""), (2, 4, ""), (3, 4, ""), (4, 5, ""), (2, 6, ""), (4, 7, ""), (5, 8, "")],
        )
    )
    diagrams.append(
        draw_flow(
            "g19_training_stages",
            "Registered training stages",
            "The staged program separates data/tokenizer gates, foundation learning, continuation, task adaptation, and identification.",
            [
                ("D0 data gate", 0.08, 0.66, TEAL),
                ("T tokenizer freeze", 0.22, 0.66, BLUE),
                ("F 300M-token foundation", 0.37, 0.66, NAVY),
                ("U 200M-token continuations", 0.53, 0.66, PURPLE),
                ("A 12M task adaptation", 0.68, 0.66, ORANGE),
                ("S 6M supervised specialization", 0.83, 0.66, RED),
                ("ID causal classifier", 0.68, 0.30, BLUE),
                ("M3B MNTP + contrastive ID", 0.83, 0.30, TEAL),
            ],
            [(0, 1, ""), (1, 2, ""), (2, 3, ""), (3, 4, ""), (4, 5, ""), (4, 6, ""), (4, 7, "")],
        )
    )
    diagrams.append(
        draw_flow(
            "g20_evaluation_tracks",
            "Evaluation tracks and evidence classes",
            "Development selection is separated from IID, source-OOD, external-transcript, romanized, robustness, and human tracks.",
            [
                ("Training", 0.10, 0.67, BLUE),
                ("Validation selection", 0.29, 0.67, TEAL),
                ("Protocol freeze", 0.48, 0.67, RED),
                ("IID test", 0.68, 0.84, NAVY),
                ("Source-OOD test", 0.68, 0.62, PURPLE),
                ("External / romanized", 0.68, 0.40, ORANGE),
                ("Robustness and routing", 0.87, 0.72, BLUE),
                ("Blinded native ratings", 0.87, 0.43, TEAL),
            ],
            [(0, 1, ""), (1, 2, "one-way"), (2, 3, ""), (2, 4, ""), (2, 5, ""), (3, 6, ""), (4, 6, ""), (5, 7, "")],
        )
    )
    diagrams.append(
        draw_flow(
            "g21_statistical_workflow",
            "Confirmatory statistical workflow",
            "Per-example paired evidence is clustered by semantic group, resampled hierarchically, tested, and multiplicity corrected.",
            [
                ("Per-example predictions", 0.10, 0.66, BLUE),
                ("Semantic-group pairing", 0.29, 0.66, TEAL),
                ("Hierarchical paired bootstrap", 0.48, 0.82, PURPLE),
                ("Paired randomization test", 0.48, 0.48, PURPLE),
                ("Registered endpoint family", 0.68, 0.66, ORANGE),
                ("Holm correction", 0.85, 0.66, RED),
                ("Claim gate", 0.85, 0.30, NAVY),
            ],
            [(0, 1, ""), (1, 2, ""), (1, 3, ""), (2, 4, ""), (3, 4, ""), (4, 5, ""), (5, 6, "CI and corrected p")],
        )
    )
    return diagrams


def build_chapters(status: dict[str, Any], tables: dict[str, pd.DataFrame]) -> list[tuple[str, str, str]]:
    main = tables["main"]
    means = main.groupby("variant").mean(numeric_only=True)
    m2_norm = means.loc["M2", "norm_macro_chrfpp"]
    m3_norm = means.loc["M3", "norm_macro_chrfpp"]
    m2_id = means.loc["M2", "id_regional_macro_f1"]
    m3_id = means.loc["M3", "id_regional_macro_f1"]
    chapters: list[tuple[str, str, str]] = []

    chapters.append(
        (
            "00_reader_guide",
            "Reader guide, scope, and evidence rules",
            f"""
This monograph documents the complete Boichitro-MoE experiment workspace as it existed at **{SNAPSHOT_LABEL}**. It is designed to remain useful after the original 91 GiB workspace is removed. The package preserves explanatory text, configurations, source-code descriptions, run metadata, reports, tables, figure source data, and a complete path-and-size inventory.

The documentation is **not a replacement for the raw experiment**. It does not contain the multi-gigabyte Parquet corpora, packed pretraining blocks, model checkpoints, optimizer state, caches, or full prediction stores. Those artifacts are indexed so that their former role and location remain auditable, but they cannot be reconstructed from this document.

Evidence is labelled conservatively:

- **Development-only** means a result was used for model, schedule, fusion, or checkpoint selection.
- **Prior-test descriptive** identifies classical or smoke outputs created before the neural protocol freeze.
- **Locked evidence** would require an immutable protocol freeze followed by scripted test execution. That evidence is not present.
- **Human evidence** requires qualified native-speaker review. The dataset review is {status['native_review_completed']}/{status['native_review_total']} complete.

The narrative distinguishes intentions from executions. The research plan and execution blueprint describe registered goals; manifests and reports describe what ran; the current filesystem snapshot determines what was present at documentation time.
""",
        )
    )
    chapters.append(
        (
            "01_executive_summary",
            "Executive summary",
            f"""
Boichitro-MoE is a provenance-first study of Bangla dialect normalization and dialect identification under source shift. The central scientific question is whether a compact, task-aware sparse decoder can learn dialect-relevant structure without using dataset-source shortcuts, while maintaining approximately matched token-active parameters against dense and sparse controls.

The data pipeline produced `boichitro_data_v1.0.0`, with 54,598 authentic normalization rows, 3,325 traceable train-only perturbations, 122,353 conflict-cleaned identification rows, 1,342 romanized source-held-out items, and 100,236 unique tokenizer-training texts. Automated engineering gates authorize internal modeling. Public redistribution and claims of linguistic validation remain blocked by the uncompleted native-speaker review.

The model family shares a 16-layer, width-512 causal decoder, grouped-query attention, RoPE, QK normalization, RMSNorm, a frozen 32k WordPiece tokenizer, and multi-token prediction. M0 is dense, M1 is Switch top-1, M2 is a shared-expert top-2 MoE, and M3 adds causal dialect evidence, task-conditioned routing, source-adversarial supervision, and GroupDRO.

At this documentation snapshot, all twelve M0–M3 development run manifests are present. This supersedes the older manuscript/readiness statement of seven completed runs. Development means are {m2_norm:.3f} versus {m3_norm:.3f} macro chrF++ for M2 and M3, and {m2_id:.4f} versus {m3_id:.4f} regional macro-F1. Thus the current development evidence does not support a claim that M3 improves on M2. Locked evaluation, registered confirmatory statistics, robustness, routing analysis, and human system ratings remain absent.

The strongest completed assets are the audited data pipeline, tokenizer freeze, dense foundation and continuation pilots, full three-seed development matrix, source-blind development fusion, systems benchmarks, 44 validated publication figures, ten reproducible paper-table families, and 107 recorded passing regression tests.
""",
        )
    )
    chapters.append(
        (
            "02_status_reconciliation",
            "Status reconciliation and chronology",
            f"""
Several artifacts were generated at different times and therefore disagree. The older Q1 readiness audit was created at `{status['old_audit_created_at']}` and recorded `{status['old_audit_main_evidence']}`. The later development snapshot records {status['main_snapshot_completed']}/{status['main_snapshot_expected']} complete runs, and the filesystem contains {status['main_manifests_present']} main run manifests.

The pipeline state file was last modified at `{status['pipeline_state_modified']}` and still reports `{status['pipeline_reported_status']}`. At documentation time, {len(status['live_training_processes_observed'])} matching live training processes were observed. The bidirectional specialist has {status['bidirectional_manifests_present']}/{status['bidirectional_expected']} completed manifests; seed 4307 contains partial training evidence but no completed run manifest.

The correct interpretation is therefore:

1. Main M0–M3 development training is complete for three seeds each.
2. The manuscript and older audit are stale with respect to those main runs.
3. The end-to-end pipeline is not complete.
4. No immutable protocol-freeze manifest or locked neural evaluation set is present.
5. The bidirectional branch is partially complete.
6. Human validation remains unstarted.

All later claims in this monograph use this reconciled chronology.
""",
        )
    )
    chapters.append(
        (
            "03_research_problem",
            "Research problem, questions, and falsifiable claim",
            """
Bangla dialect resources differ in collection source, annotation convention, orthography, and dialect coverage. When corpora are merged, dialect labels can become entangled with source identity. A classifier or normalizer may then exploit source-specific formatting and vocabulary rather than transferable dialect structure.

The primary task is source-blind dialect-to-Standard-Bangla normalization. At inference time, the system must not receive a gold dialect label or source identifier. Dialect identification is an auxiliary and independently evaluated task. General-language modeling supplies the shared foundation and replay-retention constraint.

The planned central comparison is M3 Boichitro-MoE versus M2 standard shared-expert MoE. The registered primary endpoint is source-independent normalization macro chrF++. A publishable positive claim requires a paired hierarchical confidence interval for M3−M2 entirely above zero and the associated registered randomization result to survive correction. A positive development point estimate would not be enough; the current point estimate is negative.

Supporting questions concern tokenizer fairness, router specialization, source invariance, worst-dialect behavior, replay retention, calibration, robustness, inference efficiency, and native-speaker adequacy. These outcomes cannot replace the primary endpoint.
""",
        )
    )
    chapters.append(
        (
            "04_historical_audit",
            "Historical audit and protocol repairs",
            """
The project began by auditing two local ZIP archives and two notebooks. The audit found publication-blocking problems in the earlier pipeline: a Vashantor column mismatch, held-out rows reintroduced through derived synthetic data, a missing Barishal validation file, duplicate loading of a regional corpus, template-like synthetic conflicts, uncertain ancestry inside merged BanglaDial material, model–data scale mismatch, and external diagnostics that did not support valid comparative claims.

The repaired workflow treats the original archives as immutable inputs. Every source is handled through an explicit adapter. The old derived archive is quarantined instead of silently reused. Source ancestry, licenses, dialect mapping, row origin, exclusion reason, split membership, semantic groups, and synthetic parentage are recorded in machine-readable manifests.

The monolithic notebooks are retained as historical material, not as the authoritative executable pipeline. Reusable modules under `src/boichitro`, command-line tools, YAML registries, regression tests, immutable hashes, and saved per-example outputs form the reproducible implementation.
""",
        )
    )
    chapters.append(
        (
            "05_data_provenance",
            "Data provenance, taxonomy, and canonical schema",
            """
The frozen taxonomy has thirteen labels: twelve regional varieties plus Standard Bangla. Labels are BAR, CHI, KHU, KIS, MYM, NAR, NOA, NSD, RAJ, RAN, SYL, TAN, and STD. Not every source supports every label, and the documentation preserves those missing cells rather than imputing coverage.

Canonical records include immutable row identifiers, task, source/provider, source version, source row, license record, dialect code, regional input, Standard-Bangla target when available, normalized text views, semantic component, original split provenance, frozen split, evaluation track, quality tier, synthetic flag, and ancestry.

Source adapters separate local Vashantor, regional corpora, Chatgaiyya material, Sylheti translation, pinned public datasets, external transcripts, romanized pairs, and general-language pretraining text. The data artifact ledger in this package records row counts, byte sizes, and published SHA-256 values for every frozen dataset output.
""",
        )
    )
    chapters.append(
        (
            "06_cleaning_and_leakage",
            "Cleaning, deduplication, and leakage controls",
            """
Cleaning normalizes Unicode and whitespace while preserving linguistically meaningful content. Fatal text-quality failures are excluded with explicit reasons. The pipeline applies exact pair checks, compact text checks, cross-label conflict removal, source-priority rules, and protected-evaluation ancestry controls.

Near-duplicate protection combines 64-bit SimHash with character 4-gram Jaccard filtering. Candidate pairs within a maximum Hamming distance are checked against a 0.90 Jaccard threshold. Semantic relations are converted into connected components before split assignment, preventing related rows from straddling development and protected evaluation.

The direction of removal is asymmetric by design: training relatives are removed against protected IID evaluation; source-OOD material is protected against all development inputs. A separate cross-task firewall checks that normalization evaluation inputs do not appear among identification-training inputs. Synthetic perturbations are train-only and never constitute evaluation evidence.
""",
        )
    )
    chapters.append(
        (
            "07_dataset_composition",
            "Frozen dataset composition and data gates",
            """
The final build contains 57,923 normalization rows, of which 54,598 are authentic and 3,325 are traceable train-only perturbations. The identification view contains 122,353 conflict-cleaned rows. The real romanized source-held-out track contains 1,342 rows. The tokenizer text view contains 100,236 unique texts.

The dataset report records 137,386 exclusion decisions. Large categories include same-label compact duplicates, cross-label identification conflicts, lower-priority exact normalization pairs, protected-evaluation relatives, source taxonomy exclusions, and quality failures. The full exclusion ledger is evidence of conservative filtering, not additional training data.

Automated gates pass for license resolution within the internal task scope, taxonomy, quarantine of the legacy derived archive, source-OOD protection, synthetic train-only use, compact-overlap checks, and cross-label conflict removal. The native-review gate fails because 0/230 sampled rows have been completed. Training is authorized internally; public redistribution and “linguistically validated” claims are not.
""",
        )
    )
    chapters.append(
        (
            "08_tokenizer",
            "Tokenizer study and immutable freeze",
            """
Tokenizer selection compares WordPiece, Unigram, and byte-BPE families at multiple vocabulary sizes and corpus-balance conditions. Intrinsic measures include tokens per character and byte, fertility, unknown rate, round-trip behavior, dialect dispersion, worst-to-best ratios, and Gini-style fairness summaries.

Candidates passing the intrinsic screen enter a fixed-budget proxy language-model study with seeds 1701, 2903, and 4307. Bits per character is the cross-tokenizer quality metric because token-normalized perplexity is not comparable across vocabularies. Throughput and stability provide additional practical evidence.

The selected tokenizer is `wordpiece_natural_32k`, with an actual vocabulary of 32,000. Its frozen files, metadata, and recorded hash are copied into this documentation package. Test data were not used for selection.
""",
        )
    )
    chapters.append(
        (
            "09_pretraining",
            "General Bangla pretraining corpus and dense foundation",
            """
The general-language source is a revision-pinned Bengali subset of FineWeb-2. Source verification records immutable revisions and file hashes. Filtering applies Unicode checks, Bengali-script ratios, length constraints, and direct/compact benchmark decontamination.

The recorded corpus report examined 585,619 documents and accepted 567,746. The fixed foundation budget contains 300,004,991 tokens, with a disjoint 3,000,629-token validation allocation. Packed block order and shard provenance are immutable.

The dense foundation uses sixteen decoder layers, width 512, grouped-query attention with eight query heads and two key/value heads, RoPE, QK normalization, RMSNorm, multi-token prediction, and the frozen custom tokenizer. The fixed-budget foundation run is single-seed developmental evidence.
""",
        )
    )
    chapters.append(
        (
            "10_model_family",
            "Dense, Switch, standard-MoE, and Boichitro systems",
            """
M0 is the dense control. M1 uses Switch top-1 routing and substantially more inactive expert capacity. M2 uses a shared expert plus top-2 of eight routed experts. M3 starts from the M2 continuation and adds the proposed task/dialect routing mechanisms.

The architectures are designed around approximately 83.8 million active parameters per token. Total parameters differ: about 83.8M for M0, 381.1M for M1, and 168.8M for M2/M3. This is active-parameter matching, not wall-clock or energy matching.

The first four M3 blocks remain dense. Early sparse blocks receive causal dialect-evidence signals derived only from the visible prefix. Middle layers use learned load-balance bias. Late sparse blocks receive task-conditioned bias. Shared experts remain active alongside two routed experts.
""",
        )
    )
    chapters.append(
        (
            "11_boichitro_mechanisms",
            "Boichitro routing, source adversary, and GroupDRO",
            """
The proposed contribution is not sparse capacity alone. It is the combination of causal dialect-evidence routing, late task-conditioned routing, source-adversarial supervision, and bounded group-robust optimization.

The dialect head encourages representations to preserve regional evidence without using future tokens. The gradient-reversal source head penalizes representations that predict dataset provenance. GroupDRO reweights observed dialect/source/authenticity groups within registered bounds. Router statistics track expert load, entropy, coefficient of variation, dropped-token behavior, and persistent collapse.

These mechanisms require direct validation. Current development task scores do not establish source invariance, and the locked routing-specialization outputs are absent. Therefore this document treats the mechanisms as implemented design features, not demonstrated scientific effects.
""",
        )
    )
    chapters.append(
        (
            "12_optimization_and_upcycling",
            "Optimization, continuation learning rate, and upcycling",
            """
Eligible hidden matrices use Muon, while embeddings, normalization parameters, routers, and task heads use AdamW. A separately registered AdamW-only control prevents improvements from being attributed solely to the optimizer split.

Restarting a mature foundation checkpoint at the original high learning rate produced a monotonic validation regression. A validation-only pilot selected Muon 0.001 and AdamW 0.000015 for the continuation. The rejected high-rate run remains in the archive as negative evidence.

Dense-to-MoE transfer compared abrupt bank release, unbanked transfer, random initialization, annealed cross-bank release, and permanent complementary-bank routing. Only the permanent paired-bank strategy met the registered transient and endpoint regression guards. A separate Switch pilot selected auxiliary-balance straight-through routing after a loss-free variant exhibited unacceptable load variation.
""",
        )
    )
    chapters.append(
        (
            "13_task_adaptation",
            "Task adaptation, replay retention, and identification",
            """
Task adaptation uses three seeds. Stage A consumes a fixed 12M-token mixture of general replay, dialect language modeling, normalization, and romanized material. Stage S consumes 6M tokens and selects normalization checkpoints only when general-language replay degradation remains at or below the preregistered 5% guard.

The selected `ret35_balanced` schedule allocates 30% normalization and 35% replay. It achieved 41.186 validation macro chrF++ in its pilot with 0.972% replay-NLL degradation. A higher-scoring default schedule was rejected because it incurred 15.91% degradation. This is a clear example of a protocol constraint overriding a superficially better task score.

Causal identification trains with early stopping and post-hoc temperature calibration. The separate bidirectional branch applies masked-next-token prediction and contrastive supervision before identification specialization. Two of its three seed manifests are complete at this snapshot.
""",
        )
    )
    chapters.append(
        (
            "14_baselines_and_fusion",
            "Baselines and source-blind development fusion",
            """
Normalization baselines include identity copying and a training-only word-rewrite system. The fair rewrite infers a supported dialect from source text. A legacy gold-dialect rewrite is retained strictly as an oracle diagnostic and is not a deployable comparator.

Identification controls include character TF–IDF SVM and SGD systems. These are strong IID baselines but collapse on source-OOD material, motivating source-robust modeling. External model baselines are pinned in configuration but have no completed locked manifests.

A fixed source-blind candidate selector and neural/SVM probability blend were selected on development data. They use source text, candidate outputs, and dialect probabilities inferred from source text; references, gold dialects, source IDs, and evaluation-track labels are forbidden inference features. The fixed fusion transfers without retuning across all M2/M3 development runs, but remains exploratory until locked confirmation.
""",
        )
    )
    chapters.append(
        (
            "15_evaluation_and_statistics",
            "Evaluation tracks, metrics, and statistical protocol",
            """
Normalization reports chrF++, BLEU, TER, character error rate, exact match, worst-dialect performance, and replay retention. Identification reports accuracy, balanced accuracy, thirteen-class and regional macro-F1, MCC, ECE-15, Brier score, and worst-present-dialect F1.

Tracks separate validation, group-IID test, source-held-out test, external transcript, and romanized challenge material. RAJ normalization is a zero-shot challenge and is not silently averaged into the trained-dialect endpoint.

Confirmatory inference is designed around paired per-example predictions. Semantic groups define the resampling unit. Hierarchical paired bootstrap intervals, semantic-group paired randomization, and Holm correction within registered endpoint families control dependence and multiplicity. Those locked confirmatory outputs have not been generated.
""",
        )
    )
    chapters.append(
        (
            "16_current_results",
            "Current development results and interpretation",
            f"""
All twelve main development manifests are present. Mean normalization macro chrF++ is {means.loc['M0','norm_macro_chrfpp']:.3f} for M0, {means.loc['M1','norm_macro_chrfpp']:.3f} for M1, {m2_norm:.3f} for M2, and {m3_norm:.3f} for M3. M3 trails M2 by {m3_norm-m2_norm:.3f} points in this development summary.

Mean regional identification macro-F1 is {means.loc['M0','id_regional_macro_f1']:.4f} for M0, {means.loc['M1','id_regional_macro_f1']:.4f} for M1, {m2_id:.4f} for M2, and {m3_id:.4f} for M3. M3 trails M2 by {m3_id-m2_id:.4f}. These are validation-selected outcomes and cannot support a locked-test claim.

The fixed fusion substantially improves development normalization for both M2 and M3, reaching roughly 54–55 macro chrF++ across seeds, and raises identification regional macro-F1 to roughly 0.79–0.80. M2 remains slightly stronger than M3 under this fixed transfer. The appropriate conclusion is mixed-to-negative for the proposed architectural advantage, while the source-blind system-level fusion is a promising exploratory result.
""",
        )
    )
    chapters.append(
        (
            "17_systems_efficiency",
            "Systems and inference efficiency",
            """
On the NVIDIA GB10 training benchmark, M0 processes about 15,337 tokens/s with 2.83 GiB peak memory. M1 processes about 5,211 tokens/s with 5.69 GiB. M2 and M3 process about 8,272 and 8,224 tokens/s with roughly 3.84 GiB.

M3 adds little overhead relative to M2, but both sparse models are substantially slower than dense M0 despite similar active parameter counts. Grouped expert execution, routing, dispatch, and memory movement therefore matter to practical compute claims.

Task inference benchmarks include batch-one and batched normalization and identification. The results show substantial batching gains and distinct latency/memory trade-offs. Any deployment claim should report task, batch size, generated tokens, examples per second, latency, checkpoint hash, and hardware.
""",
        )
    )
    chapters.append(
        (
            "18_reproducibility",
            "Software architecture and reproducibility",
            f"""
The project separates reusable library code under `src/boichitro`, command-line tools under `tools`, YAML configuration under `configs`, regression tests under `tests`, immutable evidence under `reports`, metrics and predictions, and run artifacts.

The recorded regression suite reports {status['recorded_tests_passed']} passed, {status['recorded_tests_failed']} failed, and one non-failing forward-compatibility warning. The 44 source figures passed 398/398 checks for paired formats, hashes, resolution, captions, and source data.

Every configuration, Python module, run manifest, training report, and CSV evidence table receives a catalog card later in this monograph. The evidence snapshot copies small human-readable artifacts so the documentation remains inspectable after the original workspace is removed.
""",
        )
    )
    chapters.append(
        (
            "19_limitations",
            "Limitations, blockers, and prohibited claims",
            """
The study is not submission-ready. No protocol-freeze manifest is present. No locked M0–M3 neural evaluation manifests, locked ablation results, registered robustness curves, locked routing specialization outputs, or confirmatory statistics are present. External baselines are configured but not completed on the locked tracks.

The native dataset review is 0/230 and blinded native system ratings are absent. Therefore the corpus cannot be described as fully linguistically validated or redistribution-ready, and machine metrics cannot establish human adequacy.

The primary M3-versus-M2 hypothesis is currently unsupported by development scores. The document must not imply a positive Boichitro effect, Q1 acceptance, or completed computational evidence. The manuscript requires author metadata, declarations, target-journal selection, full references, and updated results.
""",
        )
    )
    chapters.append(
        (
            "20_preservation_warning",
            "Preservation, deletion consequences, and recovery limits",
            """
The documentation folder preserves descriptions and small evidence, not the executable state of the full experiment. Deleting the original workspace and the complete backup will permanently remove raw and processed Parquet data, packed pretraining blocks, `.pt` checkpoints, optimizer state, caches, full predictions, and original archives unless they exist elsewhere.

GitHub normally contains the compact source repository and frozen tokenizer, not the 91 GiB research state. A PDF, DOCX, or ZIP cannot reproduce training without the omitted artifacts.

Before deleting the original or backup, retain at least one independent copy of the complete archive on an external disk or trusted object store. If deletion proceeds anyway, this documentation will preserve scientific context, tables, figures, configurations, code descriptions, hashes recorded by the experiment, run metadata, and a complete former-file inventory—but not recoverable model weights or corpora.
""",
        )
    )
    chapters.append(
        (
            "21_reproduction_guide",
            "Reproduction guide",
            """
The authoritative build order is:

1. Audit immutable local archives.
2. Build the preliminary manifest.
3. Acquire and verify pinned external sources.
4. Build the final dataset and figures.
5. Train and freeze tokenizer candidates.
6. Prepare and verify the pinned pretraining corpus and packed blocks.
7. Train the dense foundation and select continuation/upcycling/router pilots.
8. Run the three-seed M0–M3 task matrix.
9. Summarize development results and select only registered development artifacts.
10. Complete remaining bidirectional, ablation, and external studies.
11. Freeze code, configurations, data, tokenizer, checkpoints, calibration, and selection manifests.
12. Run one scripted locked evaluation, robustness, routing, statistics, and human evaluation.

The copied evidence snapshot retains the commands, configuration files, source code, reports, and run metadata needed to understand this sequence. Actual reproduction still requires the omitted data and model artifacts.
""",
        )
    )
    chapters.append(
        (
            "22_traceability",
            "Evidence traceability and documentation conventions",
            """
Each appendix card gives the original project-relative path, byte size, SHA-256 for preserved small files, structural keys, and a concise role description. Tables preview machine-readable evidence but do not alter it. Figure plates retain the original title, caption, evidence class, source-data path, and published hash.

The complete file inventory records every former project file, its category, extension, size, modification time, inode, hardlink count, and whether it was copied or merely indexed. Hardlink counts matter because checkpoints can share storage even when their logical sizes are large.

Where generated summaries conflict with older narrative text, the dated status-reconciliation chapter governs. Original reports are included unchanged as historical records and are explicitly labeled by their source timestamp.
""",
        )
    )
    for slug, title, text in chapters:
        (CHAPTERS / f"{slug}.md").write_text(f"# {title}\n\n{text.strip()}\n", encoding="utf-8")
    return chapters


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def set_cell_text(cell, text: str, bold: bool = False, size: float = 7.5, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(safe_cell(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Liberation Sans"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Bengali")
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_kv_table(doc: Document, rows: Iterable[tuple[Any, Any]], max_rows: int = 40) -> None:
    selected = list(rows)[:max_rows]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(5.2)
    table.columns[1].width = Cm(11.2)
    set_cell_text(table.rows[0].cells[0], "Field", bold=True, size=8, color="FFFFFF")
    set_cell_text(table.rows[0].cells[1], "Value", bold=True, size=8, color="FFFFFF")
    set_cell_shading(table.rows[0].cells[0], NAVY)
    set_cell_shading(table.rows[0].cells[1], NAVY)
    for key, value in selected:
        cells = table.add_row().cells
        set_cell_text(cells[0], str(key), bold=True, size=7)
        set_cell_text(cells[1], str(value), size=7)
        if len(table.rows) % 2 == 0:
            set_cell_shading(cells[0], LIGHT)
            set_cell_shading(cells[1], LIGHT)
    if len(list(rows)) > max_rows:
        paragraph = doc.add_paragraph(f"Only the first {max_rows} fields are shown here; the full file is preserved in the evidence snapshot.")
        paragraph.style = "Small Note"


def add_dataframe(
    doc: Document,
    frame: pd.DataFrame,
    max_rows: int = 14,
    max_cols: int = 8,
    font_size: float = 6.5,
) -> None:
    if frame is None or frame.empty:
        doc.add_paragraph("No rows available.")
        return
    selected = frame.iloc[:max_rows, :max_cols].copy()
    table = doc.add_table(rows=1, cols=len(selected.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, column in enumerate(selected.columns):
        set_cell_text(table.rows[0].cells[index], str(column), bold=True, size=font_size, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[index], NAVY)
    for row_index, row in selected.iterrows():
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            set_cell_text(cells[col_index], safe_cell(value, 160), size=font_size)
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[col_index], LIGHT)
    omitted_rows = max(0, len(frame) - len(selected))
    omitted_cols = max(0, len(frame.columns) - len(selected.columns))
    if omitted_rows or omitted_cols:
        paragraph = doc.add_paragraph(
            f"Preview limits: {omitted_rows} additional rows and {omitted_cols} additional columns are present in the source table."
        )
        paragraph.style = "Small Note"


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Boichitro-MoE documentation • ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("667788")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Liberation Serif"
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif Bengali")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Noto Serif Bengali")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    heading_specs = [
        ("Title", 28, NAVY),
        ("Heading 1", 20, NAVY),
        ("Heading 2", 15, BLUE),
        ("Heading 3", 12, TEAL),
        ("Heading 4", 10, PURPLE),
    ]
    for style_name, size, color in heading_specs:
        style = styles[style_name]
        style.font.name = "Liberation Sans"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Bengali")
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "DejaVu Sans Mono"
        code.font.size = Pt(7)
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Bengali")
        code.paragraph_format.left_indent = Cm(0.5)
        code.paragraph_format.right_indent = Cm(0.5)
        code.paragraph_format.space_after = Pt(4)
    if "Small Note" not in styles:
        note = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        note.font.name = "Liberation Sans"
        note.font.size = Pt(7.5)
        note.font.italic = True
        note.font.color.rgb = RGBColor.from_string("59636E")
        note._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Bengali")
        note.paragraph_format.space_after = Pt(4)
    if "Caption Custom" not in styles:
        caption = styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
        caption.font.name = "Liberation Sans"
        caption.font.size = Pt(8)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor.from_string("384B59")
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Bengali")
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Boichitro-MoE • Full Experiment Technical Documentation"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string("667788")
        add_page_number(section.footer.paragraphs[0])

    props = doc.core_properties
    props.title = "Boichitro-MoE Full Experiment Technical Documentation"
    props.subject = "Data, tokenizer, model, training, evaluation, results, code, and evidence appendices"
    props.author = "Experiment documentation generated from the Boichitro-MoE workspace"
    props.keywords = "Bangla dialect, mixture of experts, normalization, identification, reproducibility"
    props.comments = f"Documentation snapshot {SNAPSHOT_LABEL}"


def add_markdown(doc: Document, text: str, heading_offset: int = 0) -> None:
    lines = text.splitlines()
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            value = clean_markdown_inline(" ".join(item.strip() for item in paragraph_buffer))
            if value:
                doc.add_paragraph(value)
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            paragraph = doc.add_paragraph("\n".join(code_lines))
            paragraph.style = "Code Block"
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            flush_paragraph()
            level = min(4, len(heading.group(1)) + heading_offset)
            doc.add_heading(clean_markdown_inline(heading.group(2)), level=max(1, level))
            index += 1
            continue
        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1])
        ):
            flush_paragraph()
            table_lines = [stripped]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = []
            for table_line in table_lines:
                cells = [clean_markdown_inline(cell.strip()) for cell in table_line.strip("|").split("|")]
                parsed.append(cells)
            width = max(len(row) for row in parsed)
            frame = pd.DataFrame(
                [row + [""] * (width - len(row)) for row in parsed[1:]],
                columns=parsed[0] + [f"column_{i}" for i in range(len(parsed[0]), width)],
            )
            add_dataframe(doc, frame, max_rows=30, max_cols=10, font_size=6.5)
            continue
        if re.match(r"^\s*[-*]\s+", line):
            flush_paragraph()
            value = re.sub(r"^\s*[-*]\s+", "", line)
            doc.add_paragraph(clean_markdown_inline(value), style="List Bullet")
            index += 1
            continue
        if re.match(r"^\s*\d+[.)]\s+", line):
            flush_paragraph()
            value = re.sub(r"^\s*\d+[.)]\s+", "", line)
            doc.add_paragraph(clean_markdown_inline(value), style="List Number")
            index += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            paragraph = doc.add_paragraph(clean_markdown_inline(stripped.lstrip("> ")))
            paragraph.style = "Small Note"
            index += 1
            continue
        if not stripped or stripped in {"---", "***"}:
            flush_paragraph()
            index += 1
            continue
        paragraph_buffer.append(line)
        index += 1
    flush_paragraph()


def add_cover(doc: Document, status: dict[str, Any]) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BOICHITRO-MoE")
    run.font.name = "Liberation Sans"
    run.font.size = Pt(34)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY.lstrip("#"))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Full Experiment Technical Documentation")
    run.font.name = "Liberation Sans"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE.lstrip("#"))
    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.add_run(
        "Data provenance • tokenizer • pretraining • architecture • training • evaluation • "
        "development results • reproducibility • evidence appendices"
    ).font.size = Pt(11)
    for _ in range(3):
        doc.add_paragraph()
    box = doc.add_table(rows=5, cols=2)
    box.style = "Table Grid"
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_rows = [
        ("Documentation snapshot", SNAPSHOT_LABEL),
        ("Dataset version", "boichitro_data_v1.0.0"),
        ("Main development manifests", f"{status['main_manifests_present']}/12"),
        ("Native dataset review", f"{status['native_review_completed']}/{status['native_review_total']}"),
        ("Evidence status", "Development package; not locked or publication-ready"),
    ]
    for idx, (key, value) in enumerate(cover_rows):
        set_cell_text(box.rows[idx].cells[0], key, bold=True, size=9)
        set_cell_text(box.rows[idx].cells[1], value, size=9)
        if idx % 2 == 0:
            set_cell_shading(box.rows[idx].cells[0], LIGHT)
            set_cell_shading(box.rows[idx].cells[1], LIGHT)
    doc.add_paragraph()
    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning_run = warning.add_run(
        "PRESERVATION WARNING: this document does not contain the large datasets or model checkpoints."
    )
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor.from_string(RED.lstrip("#"))
    warning_run.font.size = Pt(11)
    doc.add_page_break()


def add_contents(doc: Document, chapters: list[tuple[str, str, str]]) -> None:
    doc.add_heading("Document map", level=1)
    doc.add_paragraph(
        "The PDF outline follows Word heading styles. This manual map identifies the major parts; "
        "the figure, table, configuration, code, run, and evidence indexes are also supplied as CSV files."
    )
    parts = [
        ("Part I", "Integrated technical narrative and current-status reconciliation"),
        ("Part II", "Original planning, blueprint, manuscript, and report narratives"),
        ("Part III", "Visual evidence atlas: source figures and generated diagrams"),
        ("Part IV", "Configuration cards"),
        ("Part V", "Python implementation and regression-test cards"),
        ("Part VI", "Run manifests and training-report cards"),
        ("Part VII", "CSV and JSON evidence cards"),
        ("Part VIII", "Preservation and source-file inventories"),
    ]
    add_dataframe(doc, pd.DataFrame(parts, columns=["Part", "Contents"]), max_rows=20, max_cols=2, font_size=8)
    doc.add_heading("Integrated narrative chapters", level=2)
    for number, (_, title, _) in enumerate(chapters, start=1):
        doc.add_paragraph(f"{number}. {title}", style="List Number")
    doc.add_page_break()


def add_part_title(doc: Document, part: str, title: str, description: str) -> None:
    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(part)
    run.font.name = "Liberation Sans"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL.lstrip("#"))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.font.name = "Liberation Sans"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY.lstrip("#"))
    paragraph = doc.add_paragraph(description)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style = "Small Note"


def add_figure_plate(
    doc: Document,
    number: int,
    title: str,
    image_path: Path,
    caption: str,
    metadata: list[tuple[str, Any]],
) -> None:
    doc.add_page_break()
    doc.add_heading(f"Figure {number}. {title}", level=1)
    if image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        try:
            run.add_picture(str(image_path), width=Inches(6.6))
        except Exception as exc:
            doc.add_paragraph(f"Image could not be embedded: {exc}")
    else:
        doc.add_paragraph(f"Image missing from documentation package: {image_path}")
    paragraph = doc.add_paragraph(caption)
    paragraph.style = "Caption Custom"
    add_kv_table(doc, metadata, max_rows=10)


def add_config_cards(doc: Document, configs: list[dict[str, Any]]) -> None:
    for index, item in enumerate(configs, start=1):
        doc.add_page_break()
        doc.add_heading(f"Configuration C{index:02d}: {Path(item['path']).name}", level=1)
        doc.add_paragraph(item["purpose"])
        add_kv_table(
            doc,
            [
                ("Original path", item["path"]),
                ("Bytes", item["bytes"]),
                ("SHA-256", item["sha256"]),
                ("Top-level keys", item["top_level_keys"]),
                ("Leaf settings", item["leaf_settings"]),
                ("Preserved copy", f"evidence_snapshot/{item['path']}"),
            ],
            max_rows=10,
        )
        doc.add_heading("Selected settings", level=2)
        add_kv_table(doc, flatten(item["content"]), max_rows=24)


def add_code_cards(doc: Document, code: list[dict[str, Any]]) -> None:
    for index, item in enumerate(code, start=1):
        doc.add_page_break()
        doc.add_heading(f"Implementation P{index:03d}: {item['path']}", level=1)
        doc.add_paragraph(item["role"])
        add_kv_table(
            doc,
            [
                ("Lines", item["lines"]),
                ("Bytes", item["bytes"]),
                ("SHA-256", item["sha256"]),
                ("Imports", ", ".join(item.get("imports", [])[:24])),
                ("Classes", len(item.get("classes", []))),
                ("Top-level functions", len(item.get("functions", []))),
                ("Preserved copy", f"evidence_snapshot/{item['path']}"),
            ],
            max_rows=10,
        )
        if item.get("module_docstring"):
            doc.add_heading("Module description", level=2)
            doc.add_paragraph(item["module_docstring"][:2400])
        members: list[dict[str, Any]] = []
        for value in item.get("classes", []):
            members.append(
                {
                    "kind": "class",
                    "name": value["name"],
                    "line": value["line"],
                    "signature/methods": ", ".join(value.get("methods", [])[:16]),
                    "description": value.get("docstring", "")[:240],
                }
            )
        for value in item.get("functions", []):
            members.append(
                {
                    "kind": "function",
                    "name": value["name"],
                    "line": value["line"],
                    "signature/methods": value.get("signature", "")[:200],
                    "description": value.get("docstring", "")[:240],
                }
            )
        if members:
            doc.add_heading("Public structure", level=2)
            add_dataframe(doc, pd.DataFrame(members), max_rows=22, max_cols=5, font_size=6.5)


def add_json_card(
    doc: Document,
    prefix: str,
    index: int,
    item: dict[str, Any],
    title: str,
    max_fields: int = 28,
) -> None:
    doc.add_page_break()
    doc.add_heading(f"{prefix}{index:03d}: {title}", level=1)
    if item["path"].endswith("run_manifest.json"):
        doc.add_paragraph(
            "This immutable run manifest records protocol identity, variant, seed, "
            "initialization, trainer settings, hashes, and test-access status."
        )
    elif item["path"].endswith("training_report.json"):
        doc.add_paragraph(
            "This training report records the executed stage, optimization progress, "
            "validation evidence, resource measurements, and retained checkpoint state."
        )
    else:
        doc.add_paragraph(
            "This machine-readable report is a primary evidence artifact. Its complete "
            "unaltered copy is preserved in the evidence snapshot."
        )
    add_kv_table(
        doc,
        [
            ("Original path", item["path"]),
            ("Bytes", item["bytes"]),
            ("SHA-256", item["sha256"]),
            ("Preserved copy", f"evidence_snapshot/{item['path'].replace('runs/', 'run_metadata/', 1) if item['path'].startswith('runs/') else item['path']}"),
        ],
        max_rows=8,
    )
    doc.add_heading("Structured field preview", level=2)
    add_kv_table(doc, flatten(item["content"]), max_rows=max_fields)


def add_csv_cards(doc: Document, paths: list[Path]) -> None:
    for index, path in enumerate(paths, start=1):
        doc.add_page_break()
        rel = path.relative_to(PROJECT).as_posix()
        doc.add_heading(f"Table evidence T{index:03d}: {rel}", level=1)
        try:
            frame = pd.read_csv(path)
            metadata = [
                ("Rows", len(frame)),
                ("Columns", len(frame.columns)),
                ("Bytes", path.stat().st_size),
                ("SHA-256", sha256_file(path)),
                ("Column names", ", ".join(map(str, frame.columns))),
                ("Preserved copy", f"evidence_snapshot/{rel}"),
            ]
            add_kv_table(doc, metadata, max_rows=10)
            doc.add_heading("Data preview", level=2)
            add_dataframe(doc, frame, max_rows=12, max_cols=8, font_size=6)
        except Exception as exc:
            doc.add_paragraph(f"CSV read error: {exc}")


def generate_docx(
    status: dict[str, Any],
    tables: dict[str, pd.DataFrame],
    catalogs: dict[str, Any],
    chapters: list[tuple[str, str, str]],
    generated_visuals: list[dict[str, str]],
    inventory: pd.DataFrame,
) -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc, status)
    add_contents(doc, chapters)

    add_part_title(
        doc,
        "PART I",
        "Integrated Technical Narrative",
        "Current-status synthesis grounded in the newest manifests and reports.",
    )
    for _, title, text in chapters:
        doc.add_page_break()
        doc.add_heading(title, level=1)
        add_markdown(doc, text)

    source_documents = [
        ("Experiment workspace README", PROJECT / "README.md"),
        ("Q1 execution blueprint", SOURCE_ROOT / "BANGLA_DIALECT_MOE_EXPERIMENT_BLUEPRINT.md"),
        ("Audit and research plan", SOURCE_ROOT / "BANGLA_DIALECT_MOE_Q1_RESEARCH_PLAN.md"),
        ("Consolidated results and evaluation", PROJECT / "reports/ALL_RESULTS_AND_EVALUATION.md"),
        ("Journal-neutral manuscript draft", PROJECT / "manuscript/BOICHITRO_MOE_Q1_MANUSCRIPT.md"),
        ("Final dataset report", PROJECT / "reports/FINAL_DATASET_REPORT.md"),
        ("Evaluation track contract", PROJECT / "reports/EVALUATION_TRACK_CONTRACT.md"),
        ("Local archive audit", PROJECT / "reports/LOCAL_ARCHIVE_AUDIT.md"),
        ("Q1 readiness audit", PROJECT / "reports/Q1_JOURNAL_READINESS_AUDIT.md"),
        ("Prior test-access disclosure", PROJECT / "reports/PRIOR_TEST_ACCESS_DISCLOSURE.md"),
    ]
    add_part_title(
        doc,
        "PART II",
        "Original Internal Narratives",
        "Preserved source documents. Their historical status statements are not silently updated.",
    )
    for title, path in source_documents:
        if not path.exists():
            continue
        doc.add_page_break()
        doc.add_heading(title, level=1)
        stat = path.stat()
        doc.add_paragraph(
            f"Source: {path.relative_to(SOURCE_ROOT).as_posix()} • "
            f"modified {datetime.fromtimestamp(stat.st_mtime).astimezone().isoformat()} • "
            f"SHA-256 {sha256_file(path)}"
        ).style = "Small Note"
        add_markdown(doc, read_text(path), heading_offset=1)

    add_part_title(
        doc,
        "PART III",
        "Visual Evidence Atlas",
        "All 44 validated Q1 source figures plus documentation-time diagrams and charts.",
    )
    figure_manifest = read_json(PROJECT / "figures/q1/figure_manifest.json", {})
    visual_index: list[dict[str, Any]] = []
    figure_number = 1
    for item in figure_manifest.get("figures", []):
        image_path = OUTPUT / "visuals/source_q1" / Path(item["png"]).name
        add_figure_plate(
            doc,
            figure_number,
            item.get("title", item.get("id", "")),
            image_path,
            item.get("caption", ""),
            [
                ("Figure ID", item.get("id")),
                ("Category", item.get("category")),
                ("Evidence class", item.get("evidence")),
                ("Source data", item.get("source_data")),
                ("PNG SHA-256", item.get("png_sha256")),
                ("PDF SHA-256", item.get("pdf_sha256")),
            ],
        )
        visual_index.append(
            {
                "figure_number": figure_number,
                "origin": "validated_source_figure",
                "id": item.get("id"),
                "title": item.get("title"),
                "caption": item.get("caption"),
                "path": image_path.relative_to(OUTPUT).as_posix(),
                "evidence": item.get("evidence"),
            }
        )
        figure_number += 1
    for item in generated_visuals:
        image_path = Path(item["path"])
        add_figure_plate(
            doc,
            figure_number,
            item["title"],
            image_path,
            item["caption"],
            [
                ("Figure ID", item["id"]),
                ("Origin", "Generated from documentation-time evidence"),
                ("Snapshot", SNAPSHOT_LABEL),
                ("PNG SHA-256", sha256_file(image_path)),
            ],
        )
        visual_index.append(
            {
                "figure_number": figure_number,
                "origin": "documentation_generated",
                "id": item["id"],
                "title": item["title"],
                "caption": item["caption"],
                "path": image_path.relative_to(OUTPUT).as_posix(),
                "evidence": "documentation_snapshot",
            }
        )
        figure_number += 1
    pd.DataFrame(visual_index).to_csv(APPENDICES / "visual_index.csv", index=False)

    add_part_title(
        doc,
        "PART IV",
        "Configuration Cards",
        "Every top-level YAML configuration with purpose, hash, structure, and selected settings.",
    )
    add_config_cards(doc, catalogs["configs"])

    add_part_title(
        doc,
        "PART V",
        "Implementation and Test Cards",
        "Every Python module under src, tools, and tests, described from its syntax tree.",
    )
    add_code_cards(doc, catalogs["code"])

    add_part_title(
        doc,
        "PART VI",
        "Run and Training Evidence",
        "Complete run-manifest and training-report catalogs with structured field previews.",
    )
    for index, item in enumerate(catalogs["manifests"], start=1):
        add_json_card(doc, "R", index, item, item["path"], max_fields=24)
    for index, item in enumerate(catalogs["training_reports"], start=1):
        add_json_card(doc, "TR", index, item, item["path"], max_fields=20)

    add_part_title(
        doc,
        "PART VII",
        "Machine-Readable Table and Report Evidence",
        "CSV previews and major JSON report cards; full preserved copies remain alongside this document.",
    )
    add_csv_cards(doc, catalogs["csv_paths"])
    for index, item in enumerate(catalogs["major_json"], start=1):
        add_json_card(doc, "J", index, item, item["path"], max_fields=28)

    add_part_title(
        doc,
        "PART VIII",
        "Inventories and Preservation Boundaries",
        "Summary views of all 2,577 project files and explicit identification of non-preserved large artifacts.",
    )
    doc.add_page_break()
    doc.add_heading("Complete project inventory summary", level=1)
    summary = (
        inventory.groupby("category")
        .agg(files=("path", "count"), logical_bytes=("bytes", "sum"))
        .reset_index()
        .sort_values("logical_bytes", ascending=False)
    )
    summary["logical_size"] = summary["logical_bytes"].map(human_bytes)
    add_dataframe(doc, summary, max_rows=30, max_cols=4, font_size=7)
    doc.add_paragraph(
        "The full row-level inventory is `appendices/complete_source_file_inventory.csv`. "
        "Logical bytes can double-count hard-linked checkpoints; inode and hardlink columns are retained."
    )
    doc.add_page_break()
    doc.add_heading("Large files not preserved by this documentation", level=1)
    large = inventory[inventory["bytes"] >= 100 * 1024 * 1024].sort_values("bytes", ascending=False)
    add_dataframe(doc, large[["path", "category", "human_size", "hardlink_count", "documentation_preservation"]], max_rows=40, max_cols=5, font_size=6.5)
    doc.add_paragraph(
        "Only the complete experiment backup or another independent copy can preserve these datasets and checkpoints."
    )
    doc.add_page_break()
    doc.add_heading("Final preservation notice", level=1)
    paragraph = doc.add_paragraph(
        "Do not interpret this documentation as a recoverable model/data backup. "
        "Deleting both the source workspace and complete backup permanently removes the indexed-only artifacts."
    )
    for run in paragraph.runs:
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(RED.lstrip("#"))
        run.font.size = Pt(13)
    add_kv_table(
        doc,
        [
            ("Documentation folder", str(OUTPUT)),
            ("Source workspace documented", str(PROJECT)),
            ("Snapshot", SNAPSHOT_LABEL),
            ("Source files indexed", len(inventory)),
            ("Source figure plates", len(figure_manifest.get("figures", []))),
            ("Generated visual plates", len(generated_visuals)),
            ("Configuration cards", len(catalogs["configs"])),
            ("Python cards", len(catalogs["code"])),
            ("Run manifest cards", len(catalogs["manifests"])),
            ("Training report cards", len(catalogs["training_reports"])),
            ("CSV evidence cards", len(catalogs["csv_paths"])),
            ("JSON report cards", len(catalogs["major_json"])),
        ],
        max_rows=20,
    )

    path = OUTPUT / "Boichitro_MoE_Full_Experiment_Documentation.docx"
    doc.save(path)
    return path


def write_readme(
    status: dict[str, Any],
    catalogs: dict[str, Any],
    generated_visuals: list[dict[str, str]],
    inventory: pd.DataFrame,
) -> None:
    source_figures = read_json(PROJECT / "figures/q1/figure_manifest.json", {}).get("figure_pairs", 0)
    readme = f"""# Boichitro-MoE documentation package

This is a self-contained technical documentation snapshot generated at
**{SNAPSHOT_LABEL}** from:

`{PROJECT}`

## Start here

- `Boichitro_MoE_Full_Experiment_Documentation.pdf` — primary 300+ page monograph
- `Boichitro_MoE_Full_Experiment_Documentation.docx` — editable source
- `chapters/` — integrated narrative chapters in Markdown
- `visuals/` — {source_figures} validated source figures plus {len(generated_visuals)} generated diagrams/charts
- `tables/` — generated summary tables
- `appendices/` — complete catalogs and former-file inventory
- `evidence_snapshot/` — copied configurations, code, tests, reports, tables, run metadata, and frozen tokenizer metadata

## Current evidence snapshot

- Main M0–M3 development manifests: **{status['main_manifests_present']}/12**
- Bidirectional ID manifests: **{status['bidirectional_manifests_present']}/3**
- Recorded regression tests: **{status['recorded_tests_passed']} passed, {status['recorded_tests_failed']} failed**
- Native dataset review: **{status['native_review_completed']}/{status['native_review_total']}**
- Protocol freeze and locked neural evaluation: **not complete**

## Preservation warning

This folder does **not** contain the large Parquet datasets, packed pretraining
blocks, `.pt` checkpoints, caches, or complete predictions. It contains an
index and documentation of those artifacts, not recoverable replacements.

Deleting both the original workspace and its complete backup will permanently
remove those artifacts unless another independent copy exists.

## Catalog scale

- Source files indexed: {len(inventory):,}
- YAML configuration cards: {len(catalogs['configs'])}
- Python implementation/test cards: {len(catalogs['code'])}
- Run manifest cards: {len(catalogs['manifests'])}
- Training report cards: {len(catalogs['training_reports'])}
- CSV evidence cards: {len(catalogs['csv_paths'])}
- Major JSON report cards: {len(catalogs['major_json'])}
"""
    (OUTPUT / "README.md").write_text(readme, encoding="utf-8")


def write_combined_markdown(chapters: list[tuple[str, str, str]]) -> None:
    sections = [
        "# Boichitro-MoE Full Experiment Technical Documentation",
        "",
        f"Documentation snapshot: **{SNAPSHOT_LABEL}**",
        "",
        "> This Markdown edition contains the integrated narrative. The PDF/DOCX editions also contain the full visual atlas and all evidence cards.",
        "",
    ]
    for _, title, text in chapters:
        sections.extend([f"## {title}", "", text.strip(), ""])
    (OUTPUT / "Boichitro_MoE_Integrated_Narrative.md").write_text(
        "\n".join(sections), encoding="utf-8"
    )


def write_validation_seed(
    status: dict[str, Any],
    catalogs: dict[str, Any],
    generated_visuals: list[dict[str, str]],
    inventory: pd.DataFrame,
    docx_path: Path,
) -> None:
    validation = {
        "status": "GENERATED_AWAITING_PDF_VALIDATION",
        "generated_at": GENERATED_AT.isoformat(),
        "source_root": str(PROJECT),
        "output_root": str(OUTPUT),
        "source_files_indexed": len(inventory),
        "source_figure_pairs": read_json(PROJECT / "figures/q1/figure_manifest.json", {}).get("figure_pairs", 0),
        "generated_visuals": len(generated_visuals),
        "total_visuals": read_json(PROJECT / "figures/q1/figure_manifest.json", {}).get("figure_pairs", 0)
        + len(generated_visuals),
        "configuration_cards": len(catalogs["configs"]),
        "python_cards": len(catalogs["code"]),
        "run_manifest_cards": len(catalogs["manifests"]),
        "training_report_cards": len(catalogs["training_reports"]),
        "csv_cards": len(catalogs["csv_paths"]),
        "json_cards": len(catalogs["major_json"]),
        "main_run_manifests": status["main_manifests_present"],
        "bidirectional_manifests": status["bidirectional_manifests_present"],
        "docx": {
            "path": docx_path.name,
            "bytes": docx_path.stat().st_size,
            "sha256": sha256_file(docx_path),
        },
        "pdf": None,
        "checks": {
            "output_outside_source_folder": not OUTPUT.is_relative_to(SOURCE_ROOT),
            "at_least_50_visuals": (
                read_json(PROJECT / "figures/q1/figure_manifest.json", {}).get("figure_pairs", 0)
                + len(generated_visuals)
                >= 50
            ),
            "all_12_main_manifests_documented": status["main_manifests_present"] == 12,
            "large_artifact_warning_present": True,
        },
    }
    (OUTPUT / "DOCUMENTATION_VALIDATION.json").write_text(
        json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def main() -> None:
    ensure_dirs()
    print("Indexing source files...")
    inventory = inventory_files()
    print(f"Indexed {len(inventory)} source files.")
    print("Copying small evidence snapshot...")
    copied = copy_evidence_snapshot()
    print(f"Copied {len(copied)} evidence files.")
    print("Building configuration, code, run, and table catalogs...")
    catalogs = build_catalogs()
    status = current_status()
    (APPENDICES / "documentation_time_status.json").write_text(
        json.dumps(status, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    tables = prepare_analysis_tables(status)
    print("Generating data-driven visuals...")
    generated_visuals = save_current_figures(tables, inventory)
    generated_visuals.extend(generate_flow_diagrams())
    print(f"Generated {len(generated_visuals)} new visuals.")
    chapters = build_chapters(status, tables)
    write_combined_markdown(chapters)
    print("Building DOCX monograph...")
    docx_path = generate_docx(
        status=status,
        tables=tables,
        catalogs=catalogs,
        chapters=chapters,
        generated_visuals=generated_visuals,
        inventory=inventory,
    )
    write_readme(status, catalogs, generated_visuals, inventory)
    write_validation_seed(status, catalogs, generated_visuals, inventory, docx_path)
    print(
        json.dumps(
            {
                "docx": str(docx_path),
                "docx_bytes": docx_path.stat().st_size,
                "source_files": len(inventory),
                "source_figures": read_json(PROJECT / "figures/q1/figure_manifest.json", {}).get("figure_pairs", 0),
                "generated_visuals": len(generated_visuals),
                "config_cards": len(catalogs["configs"]),
                "code_cards": len(catalogs["code"]),
                "run_manifest_cards": len(catalogs["manifests"]),
                "training_report_cards": len(catalogs["training_reports"]),
                "csv_cards": len(catalogs["csv_paths"]),
                "json_cards": len(catalogs["major_json"]),
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
